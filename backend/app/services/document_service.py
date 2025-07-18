# backend/app/services/document_service.py
"""
Document processing service for file handling, text extraction,
and vector embedding generation.
"""

import asyncio
import logging
import mimetypes
import tempfile
import os
from typing import List, Dict, Any, Optional, BinaryIO
from uuid import UUID
from datetime import datetime
from pathlib import Path

try:
    import PyMuPDF  # fitz for PDF processing
except ImportError:
    try:
        import fitz as PyMuPDF  # Alternative import name
    except ImportError:
        PyMuPDF = None
import docx
import numpy as np
import markdown
from sqlalchemy.orm import Session
import tiktoken

from app.config.settings import settings
from app.models.document import Document, DocumentChunk
from app.schemas.document import DocumentCreate
from app.db.repositories.document_repository import DocumentRepository
from app.services.file_service import FileService
from app.services.embedding_service import EmbeddingService
from app.services.vector_store import FAISSVectorStore
from app.services.simple_vector_store import SimpleVectorStore
from app.core.exceptions import DocumentProcessingError, FileStorageError

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for handling document upload, processing, and management."""
    
    def __init__(self):
        self.file_service = FileService()
        self.embedding_service = EmbeddingService()
        # Use simple vector store as default since FAISS has compatibility issues
        logger.info("Using simple vector store")
        self.vector_store = SimpleVectorStore(self.embedding_service)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # Ensure vector store directory exists
        os.makedirs(settings.vector_store_path, exist_ok=True)
    
    async def upload_document(
        self,
        file: BinaryIO,
        document_data: DocumentCreate,
        user_id: UUID,
        db: Session
    ) -> Document:
        """
        Upload document file and create database record.
        
        Args:
            file: Uploaded file object
            document_data: Document metadata
            user_id: ID of the uploading user
            db: Database session
            
        Returns:
            Created document instance
        """
        try:
            repo = DocumentRepository(db)
            
            # Upload file to storage
            file_path = await self.file_service.upload_file(
                file=file,
                bucket="documents",
                user_id=str(user_id),
                filename=document_data.filename
            )
            
            # Create document record
            document = await repo.create_document(
                document_data=document_data,
                user_id=user_id,
                file_path=file_path
            )
            
            logger.info(f"Document uploaded: {document.id}")
            return document
            
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise DocumentProcessingError(f"Failed to upload document: {str(e)}")
    
    async def process_document_content(
        self,
        document_id: UUID,
        db: Session
    ) -> None:
        """
        Process document content: extract text, create chunks, generate embeddings.
        
        Args:
            document_id: ID of document to process
            db: Database session
        """
        repo = DocumentRepository(db)
        
        try:
            # Update status to processing
            await repo.update_processing_status(document_id, "processing")
            
            # Get document
            document = await repo.get_document(document_id)
            if not document:
                raise DocumentProcessingError("Document not found")
            
            # Download file for processing
            file_content = await self.file_service.download_file(document.file_path)
            
            # Extract text based on file type
            text_content = await self._extract_text(file_content, document.file_type)
            
            if not text_content.strip():
                raise DocumentProcessingError("No text content extracted from document")
            
            # Split text into chunks
            chunks = await self._create_text_chunks(text_content, document_id)
            
            # Save chunks to database
            await repo.save_document_chunks(chunks)
            
            # Generate and save vector embeddings
            await self._generate_vector_embeddings(document_id, chunks)
            
            # Update document status
            await repo.update_processing_status(
                document_id, 
                "completed", 
                chunk_count=len(chunks)
            )
            
            logger.info(f"Document processed successfully: {document_id}")
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            await repo.update_processing_status(
                document_id, 
                "failed", 
                error_message=str(e)
            )
            raise
    
    async def _extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text content from various file types.
        
        Args:
            file_content: Raw file bytes
            file_type: MIME type of the file
            
        Returns:
            Extracted text content
        """
        try:
            if file_type == "application/pdf":
                return await self._extract_pdf_text(file_content)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_docx_text(file_content)
            elif file_type == "text/plain":
                return file_content.decode('utf-8', errors='ignore')
            elif file_type == "text/markdown":
                return await self._extract_markdown_text(file_content)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise DocumentProcessingError(f"Failed to extract text: {str(e)}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file using PyMuPDF."""
        try:
            if PyMuPDF is None:
                raise DocumentProcessingError("PyMuPDF is not installed. Please install it to process PDF files.")
            
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                text_content = ""
                pdf_doc = PyMuPDF.open(temp_file.name)
                
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                pdf_doc.close()
                return text_content
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise DocumentProcessingError(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                doc = docx.Document(temp_file.name)
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content += " | ".join(row_text) + "\n"
                
                return text_content
                
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise DocumentProcessingError(f"Failed to extract DOCX text: {str(e)}")
    
    async def _extract_markdown_text(self, file_content: bytes) -> str:
        """Extract text from Markdown file."""
        try:
            md_content = file_content.decode('utf-8', errors='ignore')
            # Convert markdown to plain text
            html = markdown.markdown(md_content)
            # Simple HTML tag removal
            import re
            clean_text = re.sub(r'<[^>]+>', '', html)
            return clean_text
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")
            raise DocumentProcessingError(f"Failed to extract Markdown text: {str(e)}")
    
    async def _create_text_chunks(
        self, 
        text_content: str, 
        document_id: UUID
    ) -> List[DocumentChunk]:
        """
        Split text into chunks and create DocumentChunk objects.
        
        Args:
            text_content: Full text content
            document_id: ID of the parent document
            
        Returns:
            List of DocumentChunk objects
        """
        try:
            # Simple text splitting based on chunk size
            chunk_size = settings.chunk_size
            chunk_overlap = settings.chunk_overlap
            
            chunks = []
            current_page = 1
            
            # Split text into chunks with overlap
            text_len = len(text_content)
            for i in range(0, text_len, chunk_size - chunk_overlap):
                chunk_text = text_content[i:i + chunk_size]
                
                if not chunk_text.strip():
                    continue
                
                # Count tokens
                token_count = len(self.tokenizer.encode(chunk_text))
                
                # Estimate page number (simple heuristic)
                if "--- Page" in chunk_text:
                    try:
                        page_marker = chunk_text.split("--- Page")[1].split("---")[0].strip()
                        current_page = int(page_marker)
                    except (IndexError, ValueError):
                        pass
                
                chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_index=len(chunks),
                    content=chunk_text,
                    token_count=token_count,
                    page_number=current_page
                )
                chunks.append(chunk)
            
            logger.info(f"Created {len(chunks)} chunks for document {document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error creating text chunks: {e}")
            raise DocumentProcessingError(f"Failed to create text chunks: {str(e)}")
    
    async def _generate_vector_embeddings(
        self, 
        document_id: UUID, 
        chunks: List[DocumentChunk]
    ) -> None:
        """
        Generate vector embeddings for document chunks and save to FAISS.
        
        Args:
            document_id: ID of the document
            chunks: List of document chunks
        """
        try:
            if not chunks:
                raise DocumentProcessingError("No chunks to process")
            
            if self.vector_store is None:
                logger.warning("Vector store not available, skipping embedding generation")
                return
            
            # Prepare texts and metadata
            texts = [chunk.content for chunk in chunks]
            metadatas = [
                {
                    "document_id": str(document_id),
                    "chunk_id": str(chunk.id),
                    "chunk_index": chunk.chunk_index,
                    "page_number": chunk.page_number,
                    "token_count": chunk.token_count,
                    "content": chunk.content
                }
                for chunk in chunks
            ]
            
            # Generate embeddings
            embeddings = await self.embedding_service.generate_embeddings(texts)
            
            # Validate embeddings format
            if not isinstance(embeddings, np.ndarray):
                logger.error(f"Embeddings must be numpy array, got {type(embeddings)}")
                raise DocumentProcessingError("Invalid embeddings format")
            
            if len(embeddings.shape) != 2:
                logger.error(f"Embeddings must be 2D array, got shape {embeddings.shape}")
                raise DocumentProcessingError("Invalid embeddings dimensions")
            
            # Create FAISS index
            await self.vector_store.create_index(document_id, embeddings, metadatas)
            
            logger.info(f"Generated embeddings for document {document_id}")
            
        except Exception as e:
            logger.error(f"Error generating vector embeddings: {e}")
            raise DocumentProcessingError(f"Failed to generate embeddings: {str(e)}")
    
    async def get_document_vector_store(self, document_id: UUID) -> bool:
        """
        Check if FAISS vector store exists for a document.
        
        Args:
            document_id: ID of the document
            
        Returns:
            True if vector store exists, False otherwise
        """
        try:
            return await self.vector_store.index_exists(document_id)
            
        except Exception as e:
            logger.error(f"Error checking vector store for document {document_id}: {e}")
            return False
    
    async def delete_document(self, document_id: UUID, db: Session) -> None:
        """
        Delete document and all associated data.
        
        Args:
            document_id: ID of document to delete
            db: Database session
        """
        try:
            repo = DocumentRepository(db)
            document = await repo.get_document(document_id)
            
            if not document:
                raise DocumentProcessingError("Document not found")
            
            # Delete file from storage
            try:
                await self.file_service.delete_file(document.file_path)
            except Exception as e:
                logger.warning(f"Error deleting file {document.file_path}: {e}")
            
            # Delete vector store
            try:
                await self.vector_store.delete_index(document_id)
            except Exception as e:
                logger.warning(f"Error deleting vector store: {e}")
            
            # Delete from database
            await repo.delete_document(document_id)
            
            logger.info(f"Document deleted: {document_id}")
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            raise DocumentProcessingError(f"Failed to delete document: {str(e)}")
    
    async def download_document_file(self, document_id: UUID) -> BinaryIO:
        """
        Download original document file.
        
        Args:
            document_id: ID of the document
            
        Returns:
            File stream
        """
        try:
            # Get document from database
            # This would need to be implemented with proper DB access
            # For now, returning a placeholder
            
            file_content = await self.file_service.download_file(f"documents/{document_id}")
            return file_content
            
        except Exception as e:
            logger.error(f"Error downloading document file: {e}")
            raise FileStorageError(f"Failed to download document: {str(e)}")
    
    async def reindex_document(self, document_id: UUID, db: Session) -> None:
        """
        Reindex document for search (regenerate embeddings).
        
        Args:
            document_id: ID of document to reindex
            db: Database session
        """
        try:
            repo = DocumentRepository(db)
            
            # Get existing chunks
            chunks, _ = await repo.get_document_chunks(document_id, 0, 1000)
            
            if not chunks:
                raise DocumentProcessingError("No chunks found for reindexing")
            
            # Regenerate vector embeddings
            await self._generate_vector_embeddings(document_id, chunks)
            
            logger.info(f"Document reindexed: {document_id}")
            
        except Exception as e:
            logger.error(f"Error reindexing document {document_id}: {e}")
            raise DocumentProcessingError(f"Failed to reindex document: {str(e)}")
    
    async def get_document_statistics(self, document_id: UUID) -> Dict[str, Any]:
        """
        Get statistics about a document.
        
        Args:
            document_id: ID of the document
            
        Returns:
            Dictionary with document statistics
        """
        try:
            stats = await self.vector_store.get_index_stats(document_id)
            return stats
            
        except Exception as e:
            logger.error(f"Error getting document statistics: {e}")
            return {"error": str(e)}
    
    async def bulk_process_documents(
        self, 
        document_ids: List[UUID], 
        db: Session
    ) -> Dict[str, Any]:
        """
        Process multiple documents in batch.
        
        Args:
            document_ids: List of document IDs to process
            db: Database session
            
        Returns:
            Processing results summary
        """
        try:
            results = {
                "successful": [],
                "failed": [],
                "total": len(document_ids)
            }
            
            # Process documents concurrently (with limit)
            semaphore = asyncio.Semaphore(settings.max_concurrent_processing)
            
            async def process_single(doc_id):
                async with semaphore:
                    try:
                        await self.process_document_content(doc_id, db)
                        results["successful"].append(str(doc_id))
                    except Exception as e:
                        results["failed"].append({
                            "document_id": str(doc_id),
                            "error": str(e)
                        })
            
            # Run all processing tasks
            await asyncio.gather(*[
                process_single(doc_id) for doc_id in document_ids
            ], return_exceptions=True)
            
            logger.info(f"Bulk processing completed: {len(results['successful'])} successful, {len(results['failed'])} failed")
            
            return results
            
        except Exception as e:
            logger.error(f"Error in bulk processing: {e}")
            raise DocumentProcessingError(f"Bulk processing failed: {str(e)}")